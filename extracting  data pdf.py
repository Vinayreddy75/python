#!/usr/bin/env python
# coding: utf-8

# In[7]:


get_ipython().system('pip install PyPDF2')


# In[8]:


import PyPDF2
from PyPDF2 import PdfFileReader


# In[11]:


PyPDF2.__version__


# In[13]:


pdf = open("file1pdf.pdf","rb")
pdf_reader = PyPDF2.PdfReader(pdf)
print("Number of pages:",len(pdf_reader.pages))
page = pdf_reader.pages[1]
print(page.extract_text())
pdf.close()


# In[1]:


import PyPDF2, urllib, nltk
from io import BytesIO
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords


# In[2]:


wFile = urllib.request.urlopen('http://www.udri.org/pdf/02%20working%20paper%201.pdf')
pdfreader = PyPDF2.PdfReader(BytesIO(wFile.read()))


# In[3]:


pageObj = pdfreader.pages[2]
page2 = pageObj.extract_text()
punctuations = ['(',')',':',':','[',']',',','...','.']
tokens = word_tokenize(page2)
stop_words = stopwords.words('english')
keywords = [word for word in tokens if not word in stop_words and not word in punctuations]


# In[4]:


keywords


# In[5]:


name_list = list()
check = ['Mr.', 'Mrs.', 'Ms.']
for idx, token in enumerate(tokens):
    if token.startswith(tuple(check)) and idx < (len(tokens) - 1):
        name = token + tokens[idx + 1] + ' ' + tokens[idx + 2]
        name_list.append(name)

print(name_list)


# In[6]:


pip install python-docx


# In[7]:


import docx


# In[9]:


doc = open("hello.docx","rb")
document = docx.Document(doc)


# In[10]:


docu=""
for para in document.paragraphs:
    docu += para.text
print(docu)


# In[11]:


for i in range(len(document.paragraphs)):

    print("The content of the paragraph "+ str(i)+" is ：" + document.paragraphs[i].text+"\n")


# In[12]:


pip install bs4


# In[13]:


import urllib.request as urllib2
from bs4 import BeautifulSoup


# In[15]:


response = urllib2.urlopen('http://www.udri.org/pdf/02%20working%20paper%201.pdf')
html_doc =response.read()


# In[ ]:


soup = BeautifulSoup(html_doc, 'html.parser')
strhtm = soup.prettify()
print(strhtm[:5000])


# In[ ]:




